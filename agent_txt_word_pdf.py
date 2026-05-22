import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
from pathlib import Path
from typing import Optional, Tuple, List
from openai import OpenAI
from dotenv import load_dotenv
import streamlit as st
import time

# 加载 .env 文件中的环境变量
load_dotenv()


class FileConverterAgent:
    """文件格式转换智能体 - 适配硅基流动平台"""

    def __init__(self, api_key: str = None, base_url: str = None):
        """
        初始化转换助手

        Args:
            api_key: 硅基流动API密钥（如果为None，则从环境变量读取）
            base_url: API基础URL
        """
        self.supported_formats = {
            'input': ['.pdf', '.docx', '.txt'],
            'output': ['.pdf', '.docx', '.txt']
        }

        # 从参数或环境变量获取配置
        if api_key is None:
            api_key = os.getenv("SILICONFLOW_API_KEY")

        if base_url is None:
            base_url = os.getenv("SILICONFLOW_BASE_URL", "https://api.siliconflow.cn/v1")

        # 初始化OpenAI客户端（硅基流动兼容OpenAI API）
        if api_key:
            self.client = OpenAI(
                api_key=api_key,
                base_url=base_url
            )
        else:
            self.client = None

    def convert_file(self, input_path: str, output_format: str) -> Tuple[str, list]:
        """
        转换文件格式

        Args:
            input_path: 输入文件路径
            output_format: 输出格式 (.pdf, .docx, .txt)

        Returns:
            输出文件路径和图片警告信息
        """
        input_path = Path(input_path)
        if not input_path.exists():
            raise FileNotFoundError(f"文件不存在: {input_path}")

        input_ext = input_path.suffix.lower()
        warnings = []

        # 检查格式支持
        if input_ext not in self.supported_formats['input']:
            raise ValueError(f"不支持的输入格式: {input_ext}")
        if output_format not in self.supported_formats['output']:
            raise ValueError(f"不支持的输出格式: {output_format}")

        # 相同格式无需转换
        if input_ext == output_format:
            return str(input_path), ["文件格式相同，无需转换"]

        # 根据输入输出格式选择转换方法
        conversion_map = {
            ('.pdf', '.docx'): self._pdf_to_word,
            ('.pdf', '.txt'): self._pdf_to_txt,
            ('.docx', '.pdf'): self._word_to_pdf,
            ('.docx', '.txt'): self._word_to_txt,
            ('.txt', '.pdf'): self._txt_to_pdf,
            ('.txt', '.docx'): self._txt_to_word,
        }

        converter = conversion_map.get((input_ext, output_format))
        if not converter:
            raise ValueError(f"不支持的转换: {input_ext} -> {output_format}")

        output_path = input_path.with_suffix(output_format)
        has_images, result_path = converter(str(input_path), str(output_path))

        # 如果有图片且转换为txt，发出警告
        if has_images and output_format == '.txt':
            warnings.append(
                "⚠️ 警告: 原文档中包含图片，转换为TXT格式后图片将会丢失！\n"
                "建议: 如需保留图片，请转换为Word或PDF格式。"
            )

        return result_path, warnings

    def analyze_image_with_multimodal(self, image_bytes: bytes, prompt: str = "请描述这张图片的内容") -> str:
        """
        使用硅基流动的多模态模型分析图片

        Args:
            image_bytes: 图片字节数据
            prompt: 提示词

        Returns:
            AI生成的图片描述
        """
        if not self.client:
            return "未配置API密钥，无法进行图片分析"

        try:
            import base64
            image_base64 = base64.b64encode(image_bytes).decode('utf-8')

            # 从环境变量获取模型名称，或使用默认值
            model = os.getenv("MULTIMODAL_MODEL", "Pro/Qwen/Qwen2-VL-7B-Instruct")

            response = self.client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{image_base64}"
                                }
                            },
                            {
                                "type": "text",
                                "text": prompt
                            }
                        ]
                    }
                ],
                max_tokens=500
            )

            return response.choices[0].message.content
        except Exception as e:
            print(f"图片分析失败: {e}")
            return f"图片分析失败: {str(e)}"

    def _check_pdf_for_images(self, pdf_path: str) -> bool:
        """检查PDF文件中是否包含图片"""
        try:
            doc = fitz.open(pdf_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images(full=True)
                if image_list:
                    doc.close()
                    return True
            doc.close()
            return False
        except Exception as e:
            print(f"检查PDF图片时出错: {e}")
            return False

    def _check_docx_for_images(self, docx_path: str) -> bool:
        """检查Word文件中是否包含图片"""
        try:
            doc = Document(docx_path)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    return True
            return False
        except Exception as e:
            print(f"检查Word图片时出错: {e}")
            return False

    def _pdf_to_word(self, pdf_path: str, word_path: str, use_multimodal: bool = False) -> Tuple[bool, str]:
        """
        PDF转Word - 提取文字和图片

        Args:
            pdf_path: PDF文件路径
            word_path: 输出Word路径
            use_multimodal: 是否使用多模态AI分析图片内容

        Returns:
            (是否有图片, 输出路径)
        """
        has_images = self._check_pdf_for_images(pdf_path)

        doc = Document()
        pdf_doc = fitz.open(pdf_path)

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]

            # 提取文本
            text = page.get_text()
            if text.strip():
                doc.add_paragraph(text)

            # 提取图片
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = pdf_doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]

                    # 添加图片到Word文档
                    image_stream = io.BytesIO(image_bytes)
                    doc.add_picture(image_stream, width=Inches(4))

                    # 可选：使用多模态AI分析图片并添加描述
                    if use_multimodal and self.client:
                        description = self.analyze_image_with_multimodal(
                            image_bytes,
                            "请用中文简要描述这张图片的内容"
                        )
                        doc.add_paragraph(f"[图片描述: {description}]")

                except Exception as e:
                    print(f"提取图片失败: {e}")

        pdf_doc.close()
        doc.save(word_path)

        return has_images, word_path

    def _pdf_to_txt(self, pdf_path: str, txt_path: str) -> Tuple[bool, str]:
        """PDF转TXT"""
        has_images = self._check_pdf_for_images(pdf_path)

        pdf_doc = fitz.open(pdf_path)
        text_content = []

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            text = page.get_text()
            if text.strip():
                text_content.append(text)
            text_content.append("\n" + "=" * 50 + "\n")  # 分页标记

        pdf_doc.close()

        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text_content))

        return has_images, txt_path

    def _word_to_pdf(self, word_path: str, pdf_path: str) -> Tuple[bool, str]:
        """Word转PDF（使用docx2pdf，Windows需要安装Microsoft Word）"""
        has_images = self._check_docx_for_images(word_path)

        try:
            from docx2pdf import convert

            # 使用docx2pdf转换
            convert(word_path, pdf_path)

            # 验证PDF文件是否生成
            if Path(pdf_path).exists():
                return has_images, pdf_path
            else:
                raise FileNotFoundError("PDF文件生成失败")

        except ImportError:
            error_msg = "缺少docx2pdf库\n\n"
            error_msg += "解决方案：\n"
            error_msg += "1. 安装docx2pdf: pip install docx2pdf\n"
            error_msg += "2. Windows系统需要安装Microsoft Word\n"
            error_msg += "3. macOS系统需要安装Microsoft Word或LibreOffice"
            raise ImportError(error_msg)
        except Exception as e:
            error_msg = f"Word转PDF失败: {str(e)}\n\n"
            error_msg += "可能的原因和解决方案：\n"
            error_msg += "1. 确保已安装docx2pdf: pip install docx2pdf\n"
            error_msg += "2. Windows: 确保安装了Microsoft Word\n"
            error_msg += "3. macOS: 确保安装了Microsoft Word或LibreOffice\n"
            error_msg += "4. Linux: 建议安装LibreOffice并使用命令行转换"
            raise RuntimeError(error_msg)

        return has_images, pdf_path

    def _word_to_txt(self, word_path: str, txt_path: str) -> Tuple[bool, str]:
        """Word转TXT"""
        has_images = self._check_docx_for_images(word_path)

        doc = Document(word_path)
        text_content = []

        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)

        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text_content))

        return has_images, txt_path

    def _txt_to_pdf(self, txt_path: str, pdf_path: str) -> Tuple[bool, str]:
        """TXT转PDF"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet

            doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            with open(txt_path, 'r', encoding='utf-8') as f:
                for line in f:
                    p = Paragraph(line.strip(), styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 6))

            doc.build(story)
        except ImportError:
            raise ImportError("需要安装reportlab库 (pip install reportlab)")

        return False, pdf_path

    def _txt_to_word(self, txt_path: str, word_path: str) -> Tuple[bool, str]:
        """TXT转Word"""
        doc = Document()

        with open(txt_path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()

        doc.save(word_path)

        return False, word_path


def main_streamlit():
    """Streamlit Web界面"""

    # 页面配置
    st.set_page_config(
        page_title="文件格式转换助手",
        page_icon="📄",
        layout="wide"
    )

    # 标题
    st.title("📄 文件格式转换助手")
    st.markdown("支持 **PDF**、**Word**、**TXT** 三种格式互转 | 基于硅基流动AI平台")

    # 侧边栏 - API状态
    with st.sidebar:
        st.header("⚙️ 设置")

        # 检查API状态
        api_configured = os.getenv("SILICONFLOW_API_KEY") is not None

        if api_configured:
            st.success("✅ API已配置")
        else:
            st.warning("⚠️ API未配置\n多模态功能不可用")
            st.info("请在 `.env` 文件中配置 `SILICONFLOW_API_KEY`")

        st.divider()

        # 使用说明
        st.header("📖 使用说明")
        st.markdown("""
        1. 上传要转换的文件
        2. 选择目标格式
        3. （可选）启用AI图片分析
        4. 点击开始转换
        5. 下载转换后的文件
        """)

        st.divider()
        st.caption("💡 提示: 含图片的文档转TXT时会丢失图片")

    # 主体区域
    col1, col2 = st.columns([2, 1])

    with col1:
        st.subheader("📤 上传文件")
        uploaded_file = st.file_uploader(
            "选择文件或拖拽到此处",
            type=['pdf', 'docx', 'txt'],
            help="支持格式: PDF, Word (.docx), TXT"
        )

        if uploaded_file:
            st.success(f"✅ 已上传: {uploaded_file.name}")
            st.info(f"文件大小: {uploaded_file.size / 1024:.2f} KB")

    with col2:
        st.subheader("⚙️ 转换选项")

        # 目标格式选择
        output_format = st.selectbox(
            "目标格式",
            options=['.docx', '.pdf', '.txt'],
            format_func=lambda x: {
                '.docx': 'Word (.docx)',
                '.pdf': 'PDF (.pdf)',
                '.txt': 'TXT (.txt)'
            }.get(x, x)
        )

        # 多模态选项
        use_multimodal = st.checkbox(
            "启用AI图片分析",
            value=False,
            help="使用多模态AI理解图片内容并生成描述（需要API密钥）",
            disabled=not api_configured
        )

        if use_multimodal and not api_configured:
            st.warning("请先配置API密钥才能使用此功能")

    # 转换按钮
    st.divider()

    if st.button("🚀 开始转换", type="primary", use_container_width=True):
        if not uploaded_file:
            st.error("❌ 请先上传文件！")
        else:
            try:
                # 创建进度条
                progress_bar = st.progress(0)
                status_text = st.empty()

                # 创建临时目录
                os.makedirs('uploads', exist_ok=True)
                os.makedirs('outputs', exist_ok=True)

                # 保存上传的文件
                input_path = os.path.join('uploads', uploaded_file.name)
                with open(input_path, 'wb') as f:
                    f.write(uploaded_file.getbuffer())

                # 显示进度
                status_text.text("🔍 正在检测文件...")
                progress_bar.progress(20)
                time.sleep(0.3)

                status_text.text("🔄 正在转换格式...")
                progress_bar.progress(50)

                # 执行转换
                agent = FileConverterAgent()
                output_file, warnings = agent.convert_file(input_path, output_format)

                progress_bar.progress(90)
                status_text.text("✅ 转换完成！")

                # 读取输出文件
                with open(output_file, 'rb') as f:
                    file_data = f.read()

                progress_bar.progress(100)
                time.sleep(0.5)

                # 显示警告信息
                if warnings:
                    st.warning("\n".join(warnings))

                # 显示成功消息
                st.success("✅ 转换成功！")

                # 下载按钮
                output_filename = Path(output_file).name
                st.download_button(
                    label="📥 下载转换后的文件",
                    data=file_data,
                    file_name=output_filename,
                    mime="application/octet-stream",
                    type="primary",
                    use_container_width=True
                )

                # 清理临时文件
                if os.path.exists(input_path):
                    os.remove(input_path)
                if os.path.exists(output_file):
                    os.remove(output_file)

            except Exception as e:
                st.error(f"❌ 转换失败: {str(e)}")
                st.exception(e)

    # 底部说明
    st.divider()
    st.markdown("""
    ### 💡 功能特点
    - ✅ 支持 PDF ↔ Word ↔ TXT 互相转换
    - ✅ 自动检测文档中的图片
    - ✅ 转TXT时提醒图片丢失风险
    - ✅ PDF转Word时保留图片
    - ✅ 可选AI图片内容分析（需API密钥）

    ### ⚠️ 注意事项
    - PDF转Word是基础的文字+图片提取，复杂排版需人工调整
    - Word转PDF需要安装docx2pdf库和Microsoft Word
    - 含有图片的文档转TXT会丢失图片
    """)


# 启动Streamlit应用
if __name__ == "__main__":
    main_streamlit()
